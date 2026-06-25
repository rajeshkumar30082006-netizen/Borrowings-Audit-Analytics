"""Report export utilities — PDF, Excel, Word, CSV."""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from utils.helpers import APP_ROOT, BRAND, format_currency


def _logo_path() -> Optional[Path]:
    logo = APP_ROOT / "assets" / "logo.png"
    return logo if logo.exists() else None


def export_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def export_excel(dfs: dict[str, pd.DataFrame], title: str = "Report") -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        for sheet, df in dfs.items():
            safe_name = sheet[:31]
            df.to_excel(writer, index=False, sheet_name=safe_name)
            workbook = writer.book
            worksheet = writer.sheets[safe_name]
            header_fmt = workbook.add_format({
                "bold": True,
                "bg_color": "#1B3A6B",
                "font_color": "white",
                "border": 1,
            })
            for col_num, _ in enumerate(df.columns):
                worksheet.write(0, col_num, df.columns[col_num], header_fmt)
    return buffer.getvalue()


def export_pdf_report(
    title: str,
    company: str,
    sections: list[dict],
    summary_kpis: Optional[dict] = None,
) -> bytes:
    """Generate branded PDF report with logo."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    navy = colors.HexColor(BRAND["navy"])
    sky = colors.HexColor(BRAND["sky"])

    title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontSize=20, textColor=navy, alignment=TA_CENTER, spaceAfter=12)
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"], fontSize=14, textColor=navy, spaceBefore=16, spaceAfter=8)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#333333"), spaceAfter=6)
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=8, textColor=colors.grey, alignment=TA_CENTER)

    story = []

    logo = _logo_path()
    if logo:
        img = Image(str(logo), width=1.2 * inch, height=1.2 * inch)
        img.hAlign = "CENTER"
        story.append(img)
        story.append(Spacer(1, 12))

    story.append(Paragraph(company, title_style))
    story.append(Paragraph(title, heading_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}", body_style))
    story.append(Spacer(1, 20))

    if summary_kpis:
        kpi_data = [["Metric", "Value"]] + [[k, str(v)] for k, v in summary_kpis.items()]
        kpi_table = Table(kpi_data, colWidths=[3 * inch, 2.5 * inch])
        kpi_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), navy),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F7FC")]),
            ("GRID", (0, 0), (-1, -1), 0.5, sky),
            ("PADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(kpi_table)
        story.append(Spacer(1, 20))

    for section in sections:
        story.append(Paragraph(section.get("title", "Section"), heading_style))
        if section.get("text"):
            story.append(Paragraph(section["text"], body_style))
        if section.get("dataframe") is not None and not section["dataframe"].empty:
            df = section["dataframe"]
            max_cols = min(len(df.columns), 6)
            cols = list(df.columns[:max_cols])
            data = [cols] + df[cols].astype(str).head(50).values.tolist()
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), navy),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ("PADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(table)
        story.append(Spacer(1, 12))

    story.append(Spacer(1, 30))
    story.append(Paragraph("CONFIDENTIAL — For audit engagement use only", footer_style))
    story.append(Paragraph(f"© {datetime.now().year} {company} · CAP AI Audit Platform", footer_style))

    doc.build(story)
    return buffer.getvalue()


def export_word_report(title: str, company: str, sections: list[dict]) -> bytes:
    doc = Document()
    logo = _logo_path()
    if logo:
        doc.add_picture(str(logo), width=Inches(1.2))
        last_para = doc.paragraphs[-1]
        last_para.alignment = 1

    heading = doc.add_heading(company, 0)
    heading.runs[0].font.color.rgb = RGBColor(27, 58, 107)
    sub = doc.add_heading(title, level=1)
    sub.runs[0].font.color.rgb = RGBColor(46, 109, 180)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")

    for section in sections:
        doc.add_heading(section.get("title", "Section"), level=2)
        if section.get("text"):
            doc.add_paragraph(section["text"])
        if section.get("dataframe") is not None and not section["dataframe"].empty:
            df = section["dataframe"].head(30)
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr[i].text = str(col)
            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(df.columns):
                    cells[i].text = str(row[col])

    doc.add_paragraph("CONFIDENTIAL — For audit engagement use only")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_executive_report(loans: pd.DataFrame, insights: list, kpis: dict, company: str) -> bytes:
    sections = [
        {"title": "Executive Summary", "text": "This report presents the borrowings and loan covenant audit findings."},
        {"title": "Loan Portfolio", "dataframe": loans},
    ]
    if insights:
        insight_df = pd.DataFrame(insights)
        sections.append({"title": "Key Observations", "dataframe": insight_df})
    return export_pdf_report("Executive Audit Report", company, sections, kpis)


def build_exception_report(errors: list[dict], company: str) -> bytes:
    df = pd.DataFrame(errors) if errors else pd.DataFrame(columns=["Rule", "Severity", "Message"])
    sections = [{"title": "Exception Report", "text": f"Total exceptions: {len(errors)}", "dataframe": df}]
    return export_pdf_report("Exception Report", company, sections)


def create_excel_template(columns: list[str], sheet_name: str = "Template") -> bytes:
    df = pd.DataFrame(columns=columns)
    return export_excel({sheet_name: df}, sheet_name)
